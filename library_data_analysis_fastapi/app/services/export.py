"""
报告导出服务
支持 Excel 和 Word 格式
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def make_docx(title: str, content: str) -> bytes:
    """生成 Word 文档"""
    doc = Document()
    
    # 标题
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 生成时间
    time_para = doc.add_paragraph()
    time_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    
    doc.add_paragraph()
    
    # 内容
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        # 处理 Markdown 标题
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("**") and "**" in line[2:]:
            # 加粗段落
            p = doc.add_paragraph()
            p.add_run(line.replace("**", "")).bold = True
        else:
            doc.add_paragraph(line)
    
    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("图书馆数据分析系统自动生成").italic = True
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_excel(data: dict, report_type: str) -> bytes:
    """生成 Excel 文件"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    
    wb = Workbook()
    ws = wb.active
    ws.title = "数据报告"
    
    # 样式
    header_font = Font(bold=True, size=12)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font_white = Font(bold=True, color="FFFFFF")
    border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    
    # 标题行
    ws.append(["指标", "数值"])
    ws["A1"].font = header_font_white
    ws["A1"].fill = header_fill
    ws["B1"].font = header_font_white
    ws["B1"].fill = header_fill
    ws.column_dimensions["A"].width = 30
    ws.column_dimensions["B"].width = 20
    
    # 数据行
    if isinstance(data, dict):
        for key, value in data.items():
            if isinstance(value, dict):
                ws.append([key, ""])
                for sub_key, sub_value in value.items():
                    ws.append([f"  {sub_key}", sub_value])
            else:
                ws.append([key, value])
    elif isinstance(data, list):
        for item in data:
            if isinstance(item, dict):
                row = [item.get(k, "") for k in item.keys()]
                ws.append(row)
            else:
                ws.append([str(item)])
    
    # 自动列宽
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            except:
                pass
        adjusted_width = min(max_length + 2, 50)
        ws.column_dimensions[column].width = adjusted_width
    
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def make_excel_for_overview(data: dict) -> bytes:
    """综合概览 Excel"""
    return make_excel(data, "overview")


def make_excel_for_reader(data: dict) -> bytes:
    """读者报告 Excel"""
    return make_excel(data, "reader")


def make_excel_for_book(data: dict) -> bytes:
    """图书报告 Excel"""
    return make_excel(data, "book")


def make_excel_for_borrow(data: dict) -> bytes:
    """借阅报告 Excel"""
    return make_excel(data, "borrow")
